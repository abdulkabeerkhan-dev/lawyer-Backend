import os
import sys
import uuid
import io
import urllib.parse
from datetime import datetime, timezone, timedelta
from typing import List, Dict, Any, Optional, cast
import re
import json

import httpx
import jwt
from jwt.algorithms import RSAAlgorithm
import sentry_sdk
from sentry_sdk.integrations.fastapi import FastApiIntegration
from pinecone import Pinecone
from anthropic import AsyncAnthropic
from supabase import create_client, Client
from dotenv import load_dotenv

from fastapi import FastAPI, HTTPException, status, Depends, Response, BackgroundTasks, Request
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

load_dotenv()

# SENTRY SYSTEM LOG ENGINE
if os.environ.get("SENTRY_DSN"):
    sentry_sdk.init(
        dsn=os.environ.get("SENTRY_DSN"),
        integrations=[FastApiIntegration()],
        traces_sample_rate=1.0,
        profiles_sample_rate=1.0,
    )

app = FastAPI(title="SECTION AI - Legal Intelligence Platform")

# CORS ORIGIN ALLOWLIST
ALLOWED_ORIGINS = [o.strip() for o in os.environ.get("ALLOWED_ORIGINS", "").split(",") if o.strip()]
if not ALLOWED_ORIGINS:
    print("⚠️ WARNING: ALLOWED_ORIGINS is not set -- CORS is mirroring ANY request origin with credentials enabled.")

def _origin_is_allowed(origin: str) -> bool:
    return (not ALLOWED_ORIGINS) or (origin in ALLOWED_ORIGINS)

@app.middleware("http")
async def dynamic_cors_middleware(request, call_next):
    origin = request.headers.get("origin")
    if request.method == "OPTIONS" and origin and _origin_is_allowed(origin):
        response = Response(status_code=200)
        response.headers["Access-Control-Allow-Origin"] = origin
        response.headers["Access-Control-Allow-Credentials"] = "true"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS, PATCH"
        req_headers = request.headers.get("access-control-request-headers")
        response.headers["Access-Control-Allow-Headers"] = req_headers or "Authorization, Content-Type, Accept, X-Requested-With, Clerk-Auth-Token"
        response.headers["Access-Control-Max-Age"] = "86400"
        return response

    response = await call_next(request)
    if origin and _origin_is_allowed(origin):
        response.headers["Access-Control-Allow-Origin"] = origin
        response.headers["Access-Control-Allow-Credentials"] = "true"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS, PATCH"
        req_headers = request.headers.get("access-control-request-headers")
        if req_headers:
            response.headers["Access-Control-Allow-Headers"] = req_headers
        else:
            response.headers["Access-Control-Allow-Headers"] = "Authorization, Content-Type, Accept, X-Requested-With, Clerk-Auth-Token"
    return response

# ENVIRONMENT CONFIGURATION
PINECONE_API_KEY = os.environ.get("PINECONE_API_KEY")
PINECONE_INDEX_NAME = os.environ.get("PINECONE_INDEX_NAME", "legal-kb-pk-local")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY")
DEV_AUTH_BYPASS_ENABLED = os.environ.get("ENABLE_DEV_AUTH_BYPASS", "false").strip().lower() == "true"
CLAUDE_MODEL = os.environ.get("CLAUDE_MODEL", "claude-haiku-4-5")

# INITIALIZE INFRASTRUCTURE CLIENTS
pinecone_index = None
if PINECONE_API_KEY:
    try:
        pc = Pinecone(api_key=PINECONE_API_KEY)
        pinecone_index = pc.Index(PINECONE_INDEX_NAME)
    except Exception as launch_err:
        print(f"⚠️ Pinecone startup warning: {launch_err}")

supabase: Any = None
if SUPABASE_URL and SUPABASE_SERVICE_KEY:
    try:
        supabase = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
    except Exception as launch_err:
        print(f"⚠️ Supabase startup warning: {launch_err}")

async_anthropic_client = None
if ANTHROPIC_API_KEY:
    try:
        raw_client = AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
        try:
            from langsmith import wrappers
            async_anthropic_client = wrappers.wrap_anthropic(raw_client)
        except Exception:
            async_anthropic_client = raw_client
    except Exception as launch_err:
        print(f"Anthropic client startup warning: {launch_err}")

security_agent = HTTPBearer(auto_error=False)
_clerk_jwks_keys_cache = None

def clean_court_name(court_name: str = "", title: str = "", case_id: str = "", text: str = "", **kwargs) -> str:
    all_fields = [str(court_name or ""), str(title or ""), str(case_id or ""), str(text or "")]
    for k, v in kwargs.items():
        if v:
            all_fields.append(str(v))
    combined = " ".join(all_fields).lower()

    if any(x in combined for x in ("ajk", "azad jammu", "azad kashmir", "mirpur", "muzaffarabad", "rawalakot")):
        if "high" in combined: return "High Court of Azad Jammu & Kashmir"
        if "service tribunal" in combined: return "AJK Service Tribunal"
        return "Supreme Court of Azad Jammu & Kashmir"

    if "federal shariat" in combined or "fsc" in combined:
        return "Federal Shariat Court"

    court_raw_lower = str(court_name or "").lower()
    case_id_lower = str(case_id or "").lower()
    is_explicit_sc = ("supreme" in court_raw_lower) or ("supreme court of pakistan" in case_id_lower) or (" pld sc " in combined)
    if is_explicit_sc:
        return "Supreme Court of Pakistan"

    if any(city in combined for city in ("lahore", "karachi", "sindh", "peshawar", "balochistan", "quetta", "islamabad")):
        if "sindh" in combined or "karachi" in combined: return "High Court of Sindh"
        if "lahore" in combined: return "Lahore High Court"
        if "peshawar" in combined: return "Peshawar High Court"
        if "balochistan" in combined or "quetta" in combined: return "High Court of Balochistan"
        if "islamabad" in combined: return "Islamabad High Court"

    if "high court" in combined:
        return "High Court"

    return str(court_name).strip().title() if court_name else "Supreme Court of Pakistan"

def format_neutral_citation(court: str, case_identifier: str, year_or_date: str) -> str:
    court_clean = clean_court_name(court)
    ident_clean = str(case_identifier).strip() if case_identifier else "Matter on Record"
    date_clean = str(year_or_date).strip() if year_or_date else ""
    
    ident_clean = re.sub(r'\b(PLD|SCMR|MLD|YLR|PCRLJ|PCrLJ|CLC|CLD|PTD)\s+\d{4}\s+[A-Za-z\s]*\d+\b', '', ident_clean, flags=re.IGNORECASE).strip()
    if not ident_clean:
        ident_clean = "Appellate Petition"

    if date_clean and date_clean not in ident_clean:
        return f"{court_clean} — {ident_clean} ({date_clean})"
    return f"{court_clean} — {ident_clean}"

def clean_markdown_formatting(text: str) -> str:
    if not text:
        return ""
    text = text.replace("**", "")
    text = re.sub(r'\[Annexure.*?\]', '', text)
    return text.strip()

def strip_control_characters(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\ufffd]', ' ', str(text))
    return re.sub(r'\s+', ' ', text).strip()

def sanitize_holding_text(text: str) -> str:
    if not text:
        return "Legal principle extracted from judgment record."
    clean_t = strip_control_characters(text)
    clean_t = re.sub(r'^\s*[\d\,\s\-\.\;\/\\]{5,}', '', clean_t).strip()
    if not clean_t or len(clean_t) < 15:
        return "Legal principle extracted from judgment record."
    digits_and_commas = len(re.findall(r'[\d\,\s]', clean_t))
    if len(clean_t) > 0 and (digits_and_commas / len(clean_t)) > 0.55:
        return "Legal principle extracted from judgment record."
    return clean_t

# AUTHENTICATION HOOKS
async def verify_clerk_session(credentials: Optional[HTTPAuthorizationCredentials] = Depends(security_agent)) -> str:
    global _clerk_jwks_keys_cache
    if not credentials:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Access Denied: Missing Authorization bearer token.")
        
    token = credentials.credentials
    clerk_secret = os.environ.get("CLERK_SECRET_KEY")

    if DEV_AUTH_BYPASS_ENABLED and (not clerk_secret or token == "mock_clerk_user_id_dev_run"):
        return "mock_clerk_user_id_dev_run"
    if not clerk_secret:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Server authentication is not configured.")

    try:
        unverified_header = jwt.get_unverified_header(token)
        if not isinstance(unverified_header, dict):
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid token signature layout shape.")
            
        kid = unverified_header.get("kid")
        if not kid:
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid token signature architecture layout.")
            
        if not _clerk_jwks_keys_cache:
            async with httpx.AsyncClient() as client:
                headers = {"Authorization": f"Bearer {clerk_secret}"}
                jwks_response = await client.get("https://api.clerk.com/v1/jwks", headers=headers)
                if jwks_response.status_code != 200:
                    raise HTTPException(status_code=status.HTTP_502_BAD_GATEWAY, detail="Failed to sync signature pairs from Clerk.")
                _clerk_jwks_keys_cache = jwks_response.json().get("keys", [])
                
        public_key = None
        if _clerk_jwks_keys_cache:
            for key_data in _clerk_jwks_keys_cache:
                if isinstance(key_data, dict) and key_data.get("kid") == kid:
                    public_key = RSAAlgorithm.from_jwk(key_data)
                    break
                
        if not public_key:
            _clerk_jwks_keys_cache = None
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Stale authentication signature validation parameters.")
            
        decoded_payload = jwt.decode(
            token,
            key=cast(Any, public_key),
            algorithms=["RS256"],
            options={"verify_aud": False},
            leeway=60
        )
        
        user_id = decoded_payload.get("sub")
        if not user_id:
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="User profile subject reference claim is missing.")
            
        return str(user_id)
        
    except jwt.exceptions.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Authentication failed: Session token has expired.")
    except Exception as error_context:
        raise HTTPException(status_code=401, detail=f"Access Denied: Token signature verification dropped: {str(error_context)}")

async def verify_admin_role(authenticated_user_id: str = Depends(verify_clerk_session)) -> str:
    if DEV_AUTH_BYPASS_ENABLED and authenticated_user_id == "mock_clerk_user_id_dev_run":
        return authenticated_user_id
        
    if not supabase:
        raise HTTPException(status_code=500, detail="Database connection is currently offline.")
    profile_query = supabase.table("users").select("role").eq("id", authenticated_user_id).execute()
    if profile_query.data and len(profile_query.data) > 0:
        first_row = profile_query.data[0]
        if isinstance(first_row, dict) and first_row.get("role") == "admin":
            return authenticated_user_id
            
    raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access Denied: Administrative permissions required.")

# DATA TRANSPORT MODELS
class UserSyncPayload(BaseModel):
    email: str
    full_name: str

class ProfileUpdatePayload(BaseModel):
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    full_name: str

class AccessRegistration(BaseModel):
    full_name: str
    firm_name: str
    email: str

class AssociateCreatePayload(BaseModel):
    full_name: str
    email: str
    status: str = "admin_approved"

class AssociateStatusPayload(BaseModel):
    status: str

class ImagePayload(BaseModel):
    image_base64: str
    image_mime_type: str

class QueryRequest(BaseModel):
    query_text: str
    images: Optional[List[ImagePayload]] = None
    category: str = "general"

class FeedbackRequest(BaseModel):
    query_id: str
    original_answer: str
    correct_answer: str

class DiaryEntryPayload(BaseModel):
    case_title: str
    case_number: str
    court_name: str
    hearing_date: str
    stage_of_case: str
    notes: Optional[str] = None

class PleadingExportRequest(BaseModel):
    court_title: str = "IN THE HIGH COURT OF SINDH AT KARACHI"
    case_title: str = "CRIMINAL / CIVIL WRIT PETITION"
    memorandum_text: str
    precedents: Optional[List[Dict[str, Any]]] = None

SYSTEM_PROMPTS = {
    "criminal": "You are an elite Pakistani criminal law specialist, holding deep expertise in the Pakistan Penal Code (PPC) and Code of Criminal Procedure (CrPC).",
    "divorce_family": "You are a leading Pakistani family law expert, specializing in the Muslim Family Laws Ordinance, Dissolution of Muslim Marriages Act, and related custody jurisprudence.",
    "government_constitutional": "You are a senior Pakistani constitutional law expert, specializing in Article 199 writ petitions, civil service regulations, and administrative law.",
    "corporate_tax": "You are a Pakistani corporate and tax law advisor, specializing in the Companies Act 2017, Contract Act, and SECP regulations.",
    "land_property": "You are an expert on Pakistani land revenue, Specific Relief Act 1877, Transfer of Property Act, and registration laws.",
    "general": "You are an elite, highly precise Pakistani legal expert and Senior Appellate Advocate."
}

def check_user_quota(user_id: str, num_images_requested: int):
    if user_id == "mock_clerk_user_id_dev_run" or not supabase:
        return
    try:
        time_limit = (datetime.now(timezone.utc) - timedelta(hours=24)).isoformat()
        res = supabase.table("queries").select("id, query_text").eq("user_id", user_id).gte("created_at", time_limit).execute()
        records = res.data if res else []
        if len(records) >= 100:
            raise HTTPException(status_code=429, detail="Daily query quota limit exceeded (Max 100 queries/day).")
            
        if num_images_requested > 0:
            vision_count = sum(1 for r in records if isinstance(r, dict) and "[Vision Context]" in str(r.get("query_text", "")))
            if vision_count >= 30:
                raise HTTPException(status_code=429, detail="Daily document upload limit exceeded (Max 30 queries with images/day).")
    except HTTPException:
        raise
    except Exception as e:
        print(f"⚠️ Quota verification error: {e}")

jobs_store: Dict[str, Dict[str, Any]] = {}

def cleanup_old_jobs():
    try:
        now = datetime.now(timezone.utc)
        expiry = timedelta(minutes=15)
        to_delete = [jid for jid, job in jobs_store.items() if now - job.get("created_at", now) > expiry]
        for jid in to_delete:
            del jobs_store[jid]
    except Exception as e:
        print(f"⚠️ Error cleaning up old jobs: {e}", file=sys.stderr)

async def process_query_job(job_id: str, request: QueryRequest, authenticated_user_id: str):
    try:
        print(f"🚀 [JOB {job_id}] Starting query execution...", file=sys.stderr, flush=True)
        
        def clean_base64_data(base64_str: str) -> str:
            return base64_str.split(",", 1)[1] if "," in base64_str else base64_str.strip()

        def sanitize_mime_type(mime: str) -> str:
            m = mime.lower().strip()
            return "image/jpeg" if m == "image/jpg" else m

        def clean_repeated_phrases(text: str) -> str:
            if not text: return ""
            text = re.sub(r'\s+', ' ', text).strip()
            prev_text = None
            while prev_text != text:
                prev_text = text
                text = re.sub(r'\b(\w+(?:\s+\w+){0,3})\s+\1\b', r'\1', text, flags=re.IGNORECASE)
            return text

        images_list = request.images or []
        check_user_quota(authenticated_user_id, num_images_requested=len(images_list))

        has_image = len(images_list) > 0
        search_keywords_query = request.query_text
        vision_input_tokens = 0
        vision_output_tokens = 0

        if has_image and async_anthropic_client and ANTHROPIC_API_KEY:
            vision_content = []
            for img in images_list:
                vision_content.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": sanitize_mime_type(img.image_mime_type),
                        "data": clean_base64_data(img.image_base64)
                    }
                })
            vision_prompt = "Identify and list 3 to 5 key legal search terms in English. Return ONLY keywords separated by spaces."
            vision_content.append({"type": "text", "text": vision_prompt})
            
            vision_message = await async_anthropic_client.messages.create(
                model=CLAUDE_MODEL,
                max_tokens=200,
                messages=[{"role": "user", "content": vision_content}]
            )
            if hasattr(vision_message, "usage") and vision_message.usage:
                vision_input_tokens = getattr(vision_message.usage, "input_tokens", 0) or 0
                vision_output_tokens = getattr(vision_message.usage, "output_tokens", 0) or 0
                
            raw_kws = "".join(getattr(b, "text", "") for b in vision_message.content).strip()
            search_keywords_query = f"{raw_kws} {request.query_text}".strip()

        mode = "simple_query"
        query_lower = request.query_text.lower()
        
        if any(k in query_lower for k in ["draft", "petition", "bail application", "plaint", "written statement", "suit for"]):
            mode = "drafting"
        elif has_image or any(k in query_lower for k in ["analyze", "contract", "fir", "agreement", "document"]):
            mode = "document_analysis"
        elif any(k in query_lower for k in ["case law", "precedent", "ruling", "judgment", "authority"]):
            mode = "caselaw_search"

        _CHITCHAT_EXACT = {
            "hi", "hello", "hey", "salam", "assalam o alaikum", "thanks", "thank you", "ok", "okay", "test", "help"
        }
        _norm_q = re.sub(r'[^\w\s]', '', request.query_text.strip().lower()).strip()
        if (not has_image) and _norm_q in _CHITCHAT_EXACT:
            chitchat_answer = "Welcome to Section AI. Ask a specific Pakistani legal proposition, cite a court petition or statute (e.g., Section 42 Specific Relief Act), or describe a case fact pattern to begin."
            if job_id in jobs_store:
                jobs_store[job_id].update({
                    "status": "done",
                    "result": {
                        "answer": chitchat_answer,
                        "citations": [],
                        "precedent_cards": [],
                        "additional_authorities": [],
                        "query_id": None,
                        "mode": "chitchat",
                        "truncated": False
                    },
                    "completed_at": datetime.now(timezone.utc),
                    "continue_state": None,
                })
            return

        COURT_ALIASES = {
            "Supreme Court of Pakistan": ["supreme court of pakistan", "supreme court"],
            "Islamabad High Court": ["islamabad high court", "ihc"],
            "Lahore High Court": ["lahore high court", "lhc"],
            "High Court of Sindh": ["sindh high court", "shc", "karachi high court", "high court of sindh"],
            "Peshawar High Court": ["peshawar high court", "phc"],
            "High Court of Balochistan": ["balochistan high court", "bhc", "quetta high court", "high court of balochistan"],
            "Federal Shariat Court": ["federal shariat court", "fsc"],
        }
        target_source = None
        for canonical, aliases in COURT_ALIASES.items():
            for alias in aliases:
                if (re.search(rf"\b{re.escape(alias)}\b", query_lower) if len(alias) <= 4 else alias in query_lower):
                    target_source = canonical
                    break
            if target_source: break

        def _expand_legal_shorthand(text: str) -> str:
            lower_text = text.lower()
            abbrev_expansions = {
                r"\bcr\.?p\.?c\.?\b": "Code of Criminal Procedure 1898 (CrPC)",
                r"\bc\.?p\.?c\.?\b": "Code of Civil Procedure 1908 (CPC)",
                r"\bp\.?p\.?c\.?\b": "Pakistan Penal Code 1860 (PPC)",
                r"\bq\.?s\.?o\.?\b": "Qanun-e-Shahadat Order 1984",
                r"\bcnsa\b": "Control of Narcotic Substances Act 1997",
                r"\bnab\b": "National Accountability Ordinance 1999",
                r"\bsra\b": "Specific Relief Act 1877",
            }
            expansions = [exp for pat, exp in abbrev_expansions.items() if re.search(pat, lower_text) and exp.lower() not in lower_text]
            expanded = re.sub(r"\bu/s\.?\s*", "under section ", text, flags=re.IGNORECASE)
            expanded = re.sub(r"\bs\.\s*(\d)", r"section \1", expanded, flags=re.IGNORECASE)
            expanded = re.sub(r"\bo\.\s*([ivxlcdm\d]+)\b", r"Order \1", expanded, flags=re.IGNORECASE)
            if expansions:
                expanded = f"{expanded} ({'; '.join(expansions)})"
            return expanded

        search_keywords_query = _expand_legal_shorthand(search_keywords_query)

        voyage_api_url = "https://api.voyageai.com/v1/embeddings"
        voyage_model = os.environ.get("VOYAGE_MODEL", "voyage-law-2")
        voyage_api_key = os.environ.get("VOYAGE_API_KEY")
        if not voyage_api_key:
            raise Exception("VOYAGE_API_KEY is missing from environment configurations!")

        async with httpx.AsyncClient(timeout=30.0) as client:
            voyage_response = await client.post(
                voyage_api_url,
                json={"input": search_keywords_query, "model": voyage_model, "input_type": "query"},
                headers={"Authorization": f"Bearer {voyage_api_key}", "Content-Type": "application/json"}
            )
            if voyage_response.status_code != 200:
                raise Exception(f"Voyage AI Embeddings API failure: {voyage_response.text}")
            query_vector = voyage_response.json()["data"][0]["embedding"]

        if not pinecone_index:
            raise Exception("Pinecone serverless engine index connection is inactive.")

        query_top_k = 60 if target_source else 30
        pinecone_kwargs = {
            "namespace": "judgments",
            "vector": query_vector,
            "top_k": query_top_k,
            "include_metadata": True
        }

        raw_matches = pinecone_index.query(**pinecone_kwargs)
        matches_list = raw_matches.get("matches", []) if isinstance(raw_matches, dict) else getattr(raw_matches, "matches", []) or []

        def format_sources_searched(retrieved_matches: List[Dict[str, Any]]) -> str:
            """
            Dynamically extracts the courts represented in the retrieved vectors.
            Prevents hardcoding 'Supreme Court of Pakistan' on High Court or provincial matters.
            """
            if not retrieved_matches:
                return "Sources Searched: Superior Courts of Pakistan"

            courts_found = set()
            for match in retrieved_matches:
                meta = match.get("metadata", {}) if isinstance(match, dict) else getattr(match, "metadata", {}) or {}
                court = meta.get("court") or meta.get("court_name")
                title = meta.get("title") or meta.get("case_title") or ""
                cid = meta.get("case_id") or ""
                cleaned = clean_court_name(str(court or "Court of Record"), title=str(title), case_id=str(cid))
                if cleaned and cleaned != "Unknown Court":
                    courts_found.add(cleaned)

            if courts_found:
                sorted_courts = sorted(list(courts_found), reverse=True)
                return "Sources Searched: " + ", ".join(sorted_courts)

            return "Sources Searched: High Courts & Supreme Court of Pakistan"

        NON_JUDGMENT_MARKERS = ["annual report", "policy document", "press release", "annual review"]
        _PAKISTANLAWSITE_RE = re.compile(r'pakistan\s*[-_]?\s*law\s*[-_]?\s*site', re.IGNORECASE)

        def _passes_source_filter(meta, target):
            normalized_court = clean_court_name(str(meta.get("court", "")), title=str(meta.get("title") or meta.get("case_title", "")), case_id=str(meta.get("case_id", "")))
            haystack = " ".join([normalized_court, str(meta.get("dataset_category", "")), str(meta.get("title", "")), str(meta.get("case_title", ""))]).lower()
            if any(marker in haystack for marker in NON_JUDGMENT_MARKERS): return False
            if any(_PAKISTANLAWSITE_RE.search(str(meta.get(k, ""))) for k in ("court", "dataset_category", "title", "case_title")): return False
            if not target: return True
            return any(alias in haystack for alias in COURT_ALIASES.get(target, [target.lower()]))

        matches_list = [m for m in matches_list if _passes_source_filter(m.get("metadata", {}) if isinstance(m, dict) else getattr(m, "metadata", {}) or {}, target_source)]

        def is_garbled_text(text: str) -> bool:
            if not text or len(text.strip()) < 10:
                return True
            if '\ufffd' in text or '\x00' in text:
                return True
            if re.search(r'\b[A-Za-z$%\\]{2,}\d+[A-Za-z$%\\]{2,}\b', text) or re.search(r'\b\d+[A-Z]{5,}\b', text):
                return True
            words = [re.sub(r'[^a-zA-Z0-9]', '', w) for w in text.split() if w.strip()]
            if not words:
                return True
            garbled_count = 0
            for w in words:
                if len(w) > 4 and sum(1 for c in w if c.isdigit()) >= 1 and sum(1 for c in w if c.isalpha()) >= 3:
                    garbled_count += 1
            if len(words) > 3 and (garbled_count / len(words)) > 0.1:
                return True
            if len(words) >= 8:
                valid_shorts = {
                    "a", "an", "the", "and", "or", "in", "on", "at", "to", "for", "of", "off", "by", "is", "it", 
                    "be", "as", "no", "not", "has", "had", "was", "per", "vs", "v", "sub", "art", "sec", "pld", 
                    "clc", "ylr", "mld", "ptd", "plc", "cld", "sc", "hc", "lhc", "shc", "phc", "bhc", "ihc", "rs", "nos",
                    "if", "do", "we", "he", "she", "me", "my", "us", "so", "up", "out", "our", "its", "may", "can", "law",
                    "act", "set", "out", "due", "any", "all", "few", "two", "one", "three", "four", "five", "six", "day"
                }
                unknown_shorts = [w for w in words if 1 <= len(w) <= 3 and w.lower() not in valid_shorts and not w.isdigit()]
                if len(unknown_shorts) / len(words) > 0.3:
                    return True
            return False

        # Case Relevance Gate: Filter political / disqualification / bar council cases for commercial & criminal queries
        POLITICAL_MARKERS = ["nawaz sharif", "imran khan", "benazir bhutto", "tikka iqbal", "zafar ali shah", "pml-n", "pti", "pakistan bar council", "bar council", "disqualification", "election petition"]
        CRIMINAL_NAB_MARKERS = ["olas khan", "national accountability ordinance", "banking companies"]
        
        is_commercial_or_criminal_query = any(k in query_lower for k in ["fir", "quash", "420", "406", "489-f", "489f", "commercial", "contract", "cheque", "bail", "specific performance", "12 sra", "banking", "recovery", "fio 2001", "leave to defend", "security deposit"])
        is_secp_or_corporate_query = any(k in query_lower for k in ["secp", "company", "companies act", "shareholder", "director", "civil court stay", "ouster of jurisdiction", "vagrancy", "ordinance 1958", "special ordinance", "12(2)", "section 12", "115 cpc", "civil revision", "42 sra", "specific relief", "fraudulent decree", "stranger", "order xxi", "order 21", "rule 97", "rule 101", "rule 103", "execution", "objection petition", "deemed decree"])

        seen_case_ids = set()
        filtered_matches = []
        for m in matches_list:
            meta = m.get("metadata", {}) if isinstance(m, dict) else getattr(m, "metadata", {}) or {}
            score = float(m.get("score", 0.0) if isinstance(m, dict) else getattr(m, "score", 0.0))
            if score < 0.45: continue

            text_content = str(meta.get("text") or meta.get("text_preview") or "").strip()
            if is_garbled_text(text_content):
                continue
            
            case_title_str = str(meta.get("title") or meta.get("case_title") or "").lower()
            if (is_commercial_or_criminal_query or is_secp_or_corporate_query) and any(pol in case_title_str for pol in POLITICAL_MARKERS):
                continue
            if is_secp_or_corporate_query and any(cr in case_title_str for cr in CRIMINAL_NAB_MARKERS):
                continue
                
            cid = meta.get("case_id") or meta.get("citation") or meta.get("title")
            if cid and cid in seen_case_ids: continue
            if cid: seen_case_ids.add(cid)
            filtered_matches.append(m)

        primary_matches = filtered_matches[:3]
        secondary_matches = filtered_matches[3:6]

        citations_payload = []
        context_parts = []

        for match in primary_matches:
            meta = match.get("metadata", {}) if isinstance(match, dict) else getattr(match, "metadata", {}) or {}
            court = clean_court_name(str(meta.get('court', 'Unknown Court')), title=str(meta.get('title', '')), case_id=str(meta.get('case_id', '')))
            year_or_date = str(meta.get('date', '') or meta.get('year', '') or 'Recent')
            case_id = str(meta.get('case_id', 'Unknown Docket'))
            text_content = str(meta.get('text', meta.get('text_preview', ''))).strip()
            title = clean_repeated_phrases(str(meta.get('title', meta.get('case_title', 'Untitled Case')) or 'Untitled Case'))
            
            neutral_cit = format_neutral_citation(court, case_id, year_or_date)
            outcome_val = str(meta.get("outcome", "")) or "Undetermined"
            statutes_val = meta.get("statutes") or []
            sections_val = meta.get("sections") or []
            match_score = float(match.get("score", 0.0) if isinstance(match, dict) else getattr(match, "score", 0.0))

            segment_text = f"CASE TITLE: {title}\nNEUTRAL CITATION: {neutral_cit}\nCOURT: {court}\nOUTCOME: {outcome_val}\nSTATUTES: {', '.join(statutes_val)}\nCONTENT: {text_content}"
            context_parts.append(segment_text)

            citations_payload.append({
                "case_id": case_id,
                "court": court,
                "year": year_or_date,
                "preview": text_content,
                "title": title,
                "citation": neutral_cit,
                "score": match_score,
                "outcome": outcome_val,
                "statutes": statutes_val,
                "sections": sections_val,
                "relevance": "High" if match_score >= 0.65 else ("Medium" if match_score >= 0.52 else "Low")
            })

        additional_authorities = []
        for match in secondary_matches:
            meta = match.get("metadata", {}) if isinstance(match, dict) else getattr(match, "metadata", {}) or {}
            court = clean_court_name(str(meta.get('court', 'Court of Record')), title=str(meta.get('title', '')), case_id=str(meta.get('case_id', '')))
            year_or_date = str(meta.get('date', '') or meta.get('year', '') or '')
            case_id = str(meta.get('case_id', ''))
            title = clean_repeated_phrases(str(meta.get('title', meta.get('case_title', 'Precedent on Record')) or 'Precedent on Record'))
            neutral_cit = format_neutral_citation(court, case_id, year_or_date)
            preview_snippet = str(meta.get('text', meta.get('text_preview', ''))).strip()[:180] + "..."
            
            additional_authorities.append({
                "title": title,
                "citation": neutral_cit,
                "summary": preview_snippet
            })

        combined_context = "\n\n=========================================\n\n".join(context_parts)

        system_prompt = """You are Section, an elite Senior Legal Research Assistant for Pakistani Appellate Advocates.

STRICT OPERATIONAL DIRECTIVES:
1. CONCISE STATUTORY SUMMARIES (NO VERBATIM TEXT DUMPS): Never quote long multi-paragraph statutory blocks. State the section number and summarize its core legal effect in one concise sentence (e.g., "Section 420 PPC penalizes cheating where fraudulent intention existed at the inception of the transaction.").
2. CRIMINAL FIR QUASHMENT DIRECTIVE (SECTION 561-A vs. ARTICLE 199):
   - Section 561-A Cr.P.C. Limitation: Inherent jurisdiction under Section 561-A applies ONLY to judicial proceedings pending in subordinate courts; it CANNOT be invoked to quash an FIR or police investigation (Shahnaz Begum, PLD 1971 SC 677; DG FIA v. Hamid Ali Shah, PLD 2023 SC 265).
   - Article 199 Writ Jurisdiction: The SOLE forum to quash an FIR during ongoing police investigation is a Constitutional Writ Petition under Article 199.
   - Trial Court Remedies (Section 249-A / 265-K): These remedies apply AFTER the Magistrate/Trial Court takes cognizance under Section 190 Cr.P.C. Their existence is a doctrine of judicial restraint, NOT an absolute jurisdictional bar to an Article 199 writ where the FIR is registered without lawful authority, is actuated by mala fides, or discloses no cognizable offence on its face.
   - Do not confuse Section 406 PPC (Criminal Breach of Trust) with Section 409 PPC (Public Servant/Banker/Agent).
3. BAIL CANCELLATION DIRECTIVE (SECTION 497(5) Cr.P.C.):
   - Proper Remedy & Forum: An application for cancellation of bail MUST be filed directly under Section 497(5) Cr.P.C. as a Criminal Miscellaneous Application (Crl. Misc.) before the Court of Session or the High Court. (NEVER cite Section 401 Cr.P.C. which deals with executive remission of sentences, and do NOT prescribe Article 199 writs where Section 497(5) is the express statutory remedy).
   - Statutory Language of Section 497(5): The High Court or Court of Session may cause any person released on bail under Section 497 to be arrested and committed to custody.
   - Grant vs. Cancellation Parameters: Considerations for grant and cancellation are entirely distinct (Tariq Bashir v. State, PLD 1995 SC 34). The High Court does NOT sit as an ordinary court of appeal to re-assess evidence.
   - Dual Grounds for Cancellation: Bail under Section 497(5) can be cancelled ONLY on two distinct grounds:
     1. Post-grant conduct/misuse: Misuse of liberty, tampering with witnesses, intimidating complainant, or flight risk; OR
     2. Patent perversity/illegality: The bail-grant order is perverse, arbitrary, fanciful, based on complete non-reading/misreading of material record, or passed without jurisdiction (Muhammad Sadiq v. Sadiq, PLD 1985 SC 182).
4. JUSTICE OF THE PEACE (SECTION 22-A / 22-B Cr.P.C.) DIRECTIVE:
   - Legal Nature: An Ex-Officio Justice of the Peace (Sessions/Additional Sessions Judge) acts as a PERSONA DESIGNATA performing ADMINISTRATIVE, EXECUTIVE, or MINISTERIAL functions, NOT judicial functions (Khizer Hayat v. IGP Punjab, PLD 2005 Lah 470 [Full Bench]; Muhammad Bashir v. SHO, PLD 2007 SC 893; PLD 2014 SC 753).
   - Incompetency of Statutory Revisions: Neither a Criminal Revision (Section 439/435 Cr.P.C.) nor a Civil Revision (Section 115 CPC) lies against an order passed under Section 22-A(6) Cr.P.C. (NEVER claim a Criminal Revision or Civil Revision lies).
   - Maintainable Remedy: Because no statutory appeal or revision is provided by law against a Section 22-A order, the SOLE and EXCLUSIVE remedy for an aggrieved party is a CONSTITUTIONAL WRIT PETITION under Article 199 of the Constitution of Pakistan 1973 before the High Court.
   - Standard of Interference under Article 199: High Court interferes only if the order is without jurisdiction, coram non judice, or passed without verifying whether the applicant first approached the SHO/SP under Section 154 Cr.P.C.
4. SECTION 12(2) CPC & CIVIL REVISION DIRECTIVE:
   - Mandatory Bar: Section 12(2) CPC bars an independent suit under Section 42 SRA for PARTIES to the suit and their privies / representatives-in-interest.
   - The Stranger Exception: A third-party stranger whose independent title or rights are affected by a collusive/fraudulent decree between other parties MAY institute an independent regular suit for declaration under Section 42 SRA (Mst. Safia Bibi v. Mst. Aisha Bibi; Mst. Sughran Bibi v. Aziz Begum; Messrs Crescent Petroleum v. M.V. Monchegorsk).
   - Revisional Remedy & Forum: If a trial court dismisses a Section 12(2) application, the remedy is a CIVIL REVISION under Section 115 CPC (heard by a Single Bench, first before the District Court / Additional District Judge, then High Court). It is NEVER a 'Criminal Revision' and NEVER before a Division Bench.
   - Limitation: Governed by Article 181 of the Limitation Act 1908 (3 years from the date when right to apply accrues / discovery of fraud under Section 18 Limitation Act), NOT an internal proviso within Section 12(2) (Section 12(2) CPC contains no internal proviso).
5. SECTION 115 CPC & ORDER XXXIX INJUNCTION DIRECTIVE:
   - Statutory Standard: Revisional interference under Section 115 CPC is strictly confined to Clauses (a), (b), and (c):
     (a) exercise of jurisdiction not vested;
     (b) failure to exercise jurisdiction vested;
     (c) acting in the exercise of jurisdiction ILLEGALLY or with MATERIAL IRREGULARITY.
     (DO NOT use criminal revision terminology like 'illegal, unjust, or improper').
   - Interlocutory Discipline: Discretion exercised under Order XXXIX Rules 1 & 2 CPC cannot be disturbed in Civil Revision unless tainted by patent perversity, misreading/non-reading of material record, or jurisdictional defect (Muhammad Umar Beg v. Sultan Mahmood, PLD 1970 SC 139).
   - Conjunctive Triad: Prima facie case, balance of convenience, and irreparable loss are CONJUNCTIVE. If concurrent findings establish no prima facie case, relief is precluded; the High Court cannot grant an injunction or blanket status quo in revision without first setting aside the factual finding on grounds of material irregularity.
6. ORDER XXI RULES 97–103 CPC (EXECUTION & THIRD-PARTY OBJECTIONS) DIRECTIVE:
   - Mandatory Inquiry: Under Order XXI Rule 97, an independent third party asserting bona fide title/possession CANNOT be summarily dispossessed by police force. The executing court must treat the objection as a lis, frame issues, and record evidence.
   - Express Statutory Bar on Separate Suits (Rule 101): Under Order XXI Rule 101 CPC (as substituted by LRO 1972), all questions of title, right, or interest arising between the parties or third-party objectors SHALL be determined by the executing court and NOT BY A SEPARATE SUIT. An independent suit under Section 42 SRA is strictly barred.
   - Deemed Decree & Appellate Remedy (Rule 103): Under Order XXI Rule 103 CPC (as substituted by LRO 1972), an order adjudicating an application under Rule 98, 99, or 101 HAS THE FORCE OF A DECREE and is subject to the same conditions as to appeal. It is APPEALABLE as a Regular First Appeal (RFA) under Section 96 CPC, NOT a mere revision or interlocutory order.
7. BANK ACCOUNT FREEZE / BLOCK DIRECTIVE (SECTION 550 Cr.P.C. / AMLA 2010 / NAB / SBP):
   - Police / Investigative Freezes under Section 550 Cr.P.C.:
     * Mandatory Procedure: Freezing of a bank account under Section 550 Cr.P.C. by police / FIA is a SEIZURE of property. The investigating officer MUST immediately report the seizure to the Area Magistrate.
     * Direct Nexus Required: An account can ONLY be frozen if there is a direct, evidence-backed nexus between the specific funds in the account and the alleged offence. Blanket freezing of personal or business accounts without nexus is ILLEGAL (Bank Alfalah v. FIA; Habib Bank Ltd. v. State).
     * Area Magistrate Remedy: The account holder can apply directly under Section 516-A / 523 Cr.P.C. before the Area Magistrate for unfreezing / de-blocking.
   - Anti-Money Laundering Act 2010 (AMLA) Freezes (Sections 8, 12, & 19):
     * Provisional freeze by FMU / investigating agency requires prior written approval from the High Court / Special Court. Indefinite administrative freezes without court confirmation are ultra vires.
   - NAB Freezes under Section 12 NAO 1999:
     * Requires order of Chairman NAB and confirmation by Accountability Court within 30 days.
   - Maintainable Remedy: Where a bank account is frozen arbitrarily, without statutory notice, or without Magistrate report, a CONSTITUTIONAL WRIT PETITION under Article 199 lies directly before the High Court for de-blocking / unfreezing.
   - NEVER state that the database lacks authorities or that bank account freezes cannot be answered; cite Section 550 Cr.P.C., Section 516-A/523 Cr.P.C., AMLA 2010, and Article 199 writ jurisdiction authoritatively.
8. STATUTORY FACTUAL INTEGRITY RULE:
   - When asked for a specific section of a statute or special ordinance, if the exact text of that statute is NOT present in the retrieved database context, state clearly: "The exact statutory text of [Statute Name] is not indexed in the verified database."
   - DO NOT fabricate section numbers, definitions, or 19th-century British common-law classifications (e.g., 'rogues and vagabonds' or 'idle and disorderly persons') to fill gaps.
   - DO NOT substitute specialized criminal statutes (e.g., NAB Ordinance 1999 or Banking Companies Ordinance 2001) when queried on provincial or local ordinances (such as the West Pakistan Vagrancy Ordinance 1958, W.P. Ordinance XX of 1958) unless corruption or NAB accountability is explicitly raised.
   - If an offence under a special statute is bailable or cognizable (e.g., Section 19 of the West Pakistan Vagrancy Ordinance 1958), note that bailable offences are governed by Section 496 Cr.P.C. before the Area Magistrate / Police Station in-charge as a matter of statutory right (NOT Article 199 writ petitions).
8. CORPORATE & REGULATORY JURISDICTION DIRECTIVE:
   - When queried on corporate disputes, SECP, or company affairs:
     a. Rely primarily on the Companies Act 2017 (specifically Section 5 for High Court Company Bench jurisdiction, and Section 481 for the express ouster of civil court jurisdiction) and the Securities and Exchange Commission of Pakistan Act 1997 (Act XLII of 1997).
     b. Ground civil court jurisdiction analysis in Section 9 of the Code of Civil Procedure (CPC 1908) regarding express or implied statutory bars, and Section 10 CPC (Stay of suits / res sub judice).
     c. Do NOT substitute specialized criminal statutes (e.g., NAB Ordinance 1999 or Banking Companies Ordinance 2001) for corporate/commercial jurisdiction unless criminal liability is explicitly raised.
     d. Correct Statute Naming: The governing SECP statute is the "Securities and Exchange Commission of Pakistan Act 1997 (Act XLII of 1997)" (NEVER call it "SECP Act 2017"). The governing corporate statute is the "Companies Act 2017".
9. SOURCE ATTRIBUTION DIRECTIVE:
   - Do NOT hardcode 'Sources Searched: Supreme Court of Pakistan' into body text or headers.
   - If referencing sources, dynamically state the exact forum(s) involved in the cited authorities (e.g., 'Lahore High Court', 'High Court of Sindh', 'Supreme Court of Pakistan').
   - When addressing High Court writs (Article 199), Revisions (Section 115 CPC), or Intra-Court Appeals (Section 3 Law Reforms Ordinance 1972), explicitly acknowledge High Court Division Bench jurisprudence alongside Supreme Court precedents.
10. CASE RELEVANCE GATE: Cite ONLY top 2-3 precedents where the ratio directly governs the subject matter. Disregard political or constitutional disqualification cases.
11. BUDGETING & COMPLETION RULE:
   - Ensure the output NEVER terminates mid-sentence.
   - Balance length across all sections so Section IV (Procedural & Strategic Litigation Playbook) and the Prayer Clause Formulation are ALWAYS fully articulated before proceeding to Sources & Provenance.

OUTPUT STRUCTURE:
<<<CARDS>>>
[
  {
    "case_name": "Party Names",
    "case_id": "Canonical Case ID Slug",
    "citation": "Official Court and Petition / Docket Number",
    "date": "Year or exact date",
    "issue": "Detailed legal question resolved.",
    "holding": "Two concise sentences on the ratio.",
    "why_relevant": "One sentence applying directly to the user's issue.",
    "statutes_invoked": [{"name": "Statute Name and Section", "explanation": "Statutory function & mandate"}],
    "outcome": "1-2 words (e.g. 'Bail Allowed', 'Dismissed')",
    "verified_source": true
  }
]
<<<END_CARDS>>>

<<<ANSWER>>>
### I. EXECUTIVE SUMMARY & LEGAL OPINION
[Formal Senior Advocate executive conclusion and core legal opinion].

---

### II. CONTROLLING STATUTORY ARCHITECTURE

[Statute Name & Section]
> *"[Verbatim exact statutory provision or precise statutory quote in italics]"*

* Legislative Intent & Doctrinal Scope: [Explanation of legislative intent and statutory scope]
* Procedural Application to Instant Case: [Direct application to the user's scenario]

---

### III. CONTROLLING JUDICIAL PRECEDENTS & APPELLATE RATIO

[Party Names] ([Citation]) — [Court Name]
> *"[Verbatim core holding or key ratio quote from judgment]"*

* Factum & Doctrinal Controversy: [Factual dispute and controversy]
* Appellate Holding & Principle: [Appellate holding and governing principle]
* Direct Applicability & Distinguishing Factors: [Direct application to the instant dispute]

---

### IV. PROCEDURAL & STRATEGIC LITIGATION PLAYBOOK

* Form of Pleading & Proper Forum: [Proper forum, bench, and pleading form]
* Jurisdictional Threshold & Gatekeeping: [Threshold requirements and gatekeeping standards]
* Limitation & Pre-Condition for Relief: [Limitation periods and conditional stay requirements]
<<<END_ANSWER>>>

CONSTRAINTS:
- MANDATORY 4-PART FORMAT: YOU MUST FOLLOW THIS EXACT 4-PART FORMAT FOR EVERY SINGLE RESPONSE NO MATTER WHAT:
  `### I. EXECUTIVE SUMMARY & LEGAL OPINION`
  `---`
  `### II. CONTROLLING STATUTORY ARCHITECTURE`
  `---`
  `### III. CONTROLLING JUDICIAL PRECEDENTS & APPELLATE RATIO`
  `---`
  `### IV. PROCEDURAL & STRATEGIC LITIGATION PLAYBOOK`
- NEVER write custom section titles like "### BANK ACCOUNT FREEZE/BLOCK JURISPRUDENCE".
- NEVER state "I cannot synthesize case law" or "I must advise that the database context does not contain reported judgments". Formulate a complete Senior Advocate legal opinion using Section 550 Cr.P.C., Section 516-A/523 Cr.P.C., AMLA 2010, NAO 1999, and Article 199 writ jurisdiction authoritatively inside the 4-part legal opinion layout.
- ABSOLUTELY NO DOUBLE ASTERISKS (**): DO NOT write double asterisks (**) anywhere in the response text under any circumstances. All titles, statute names, and bullet headers MUST be written as plain text without any ** asterisks.
- Follow the NAME -> EXPLAIN -> APPLY legal reasoning structure for every statutory provision and judicial precedent.
- NEVER truncate mid-sentence. Budget output length cleanly.
"""

        claude_user_message = f"Context from Legal Database:\n{combined_context}\n\nQuestion: {request.query_text}"

        final_kwargs = {
            "model": CLAUDE_MODEL,
            "max_tokens": 8192,
            "system": system_prompt,
            "messages": [{"role": "user", "content": claude_user_message}]
        }

        claude_message = await async_anthropic_client.messages.create(**final_kwargs)
        raw_model_output = "".join(getattr(b, "text", "") for b in claude_message.content).strip()
        is_token_truncated = (getattr(claude_message, "stop_reason", None) == "max_tokens")

        executive_answer = ""
        precedent_cards = []

        cards_match = re.search(r'<<<CARDS>>>(.*?)<<<END_CARDS>>>', raw_model_output, re.DOTALL)
        if cards_match:
            try:
                cards_str = cards_match.group(1).strip()
                parsed_cards = json.loads(cards_str)
                if isinstance(parsed_cards, list):
                    precedent_cards = parsed_cards
            except Exception:
                card_objs = re.findall(r'\{\s*"case_name".*?\}', cards_match.group(1), re.DOTALL)
                for c_str in card_objs:
                    try:
                        c_json = json.loads(c_str)
                        if isinstance(c_json, dict) and "case_name" in c_json:
                            precedent_cards.append(c_json)
                    except Exception:
                        pass

        answer_match = re.search(r'<<<ANSWER>>>(.*?)<<<END_ANSWER>>>', raw_model_output, re.DOTALL)
        if answer_match:
            executive_answer = answer_match.group(1).strip()
        else:
            executive_answer = re.sub(r'<<<CARDS>>>.*?<<<END_CARDS>>>', '', raw_model_output, flags=re.DOTALL)
        executive_answer = clean_markdown_formatting(executive_answer)

        for idx, card in enumerate(precedent_cards):
            if idx < len(citations_payload):
                card["raw_judgment_text"] = strip_control_characters(citations_payload[idx].get("preview", ""))
                card["citation"] = citations_payload[idx].get("citation", card.get("citation"))
                card["case_id"] = citations_payload[idx].get("case_id")
            card["holding"] = sanitize_holding_text(card.get("holding", ""))

        if not precedent_cards and citations_payload:
            precedent_cards = [
                {
                    "case_name": c["title"],
                    "case_id": c["case_id"],
                    "citation": c["citation"],
                    "date": c["year"],
                    "issue": "Legal proposition extracted from indexed public judgment record.",
                    "holding": sanitize_holding_text(c.get("preview", "")[:250]),
                    "why_relevant": "Retrieved precedent directly governing the statutory issues raised.",
                    "statutes_invoked": [{"name": s, "explanation": "Governing statutory authority"} for s in c.get("statutes", [])],
                    "outcome": c.get("outcome", "Undetermined"),
                    "verified_source": True,
                    "raw_judgment_text": strip_control_characters(c.get("preview", ""))
                }
                for c in citations_payload
            ]

        add_authorities_text = ""
        if additional_authorities:
            add_authorities_lines = ["\n\nADDITIONAL RELEVANT AUTHORITIES:"]
            for auth in additional_authorities:
                add_authorities_lines.append(f"• {auth['title']} — {auth['citation']}")
            add_authorities_text = "\n".join(add_authorities_lines)

        dynamic_footer = format_sources_searched(primary_matches + secondary_matches)
        display_answer = executive_answer + add_authorities_text + "\n\n" + dynamic_footer

        inserted_row_id = str(uuid.uuid4())
        if supabase:
            db_insert = supabase.table("queries").insert({
                "user_id": authenticated_user_id,
                "query_text": f"[Vision Context] {request.query_text}" if has_image else request.query_text,
                "answer_text": display_answer,
                "citations": citations_payload,
                "input_tokens": getattr(claude_message.usage, "input_tokens", 0) + vision_input_tokens,
                "output_tokens": getattr(claude_message.usage, "output_tokens", 0) + vision_output_tokens
            }).execute()
            if db_insert.data and len(db_insert.data) > 0:
                inserted_row_id = str(db_insert.data[0].get("id", inserted_row_id))

        if job_id in jobs_store:
            jobs_store[job_id].update({
                "status": "done",
                "result": {
                    "answer": display_answer,
                    "precedent_cards": precedent_cards,
                    "additional_authorities": additional_authorities,
                    "citations": citations_payload,
                    "query_id": inserted_row_id,
                    "mode": mode,
                    "truncated": is_token_truncated
                },
                "completed_at": datetime.now(timezone.utc),
                "continue_state": {
                    "system_prompt": system_prompt,
                    "claude_message_content": claude_user_message,
                    "raw_model_answer": display_answer,
                    "precedent_cards": precedent_cards,
                    "citations_payload": citations_payload,
                    "mode": mode,
                    "category": request.category,
                    "target_source": target_source,
                    "inserted_row_id": inserted_row_id,
                    "continuation_rounds": 0,
                }
            })

    except Exception as e:
        import traceback
        traceback.print_exc(file=sys.stderr)
        if job_id in jobs_store:
            jobs_store[job_id].update({
                "status": "error",
                "error": str(e),
                "completed_at": datetime.now(timezone.utc)
            })

# FULL JUDGMENT RETRIEVAL WITH PATH-SAFE DOCKET/CITATION PARSING & REASSEMBLY
@app.get("/judgment/{case_id:path}")
async def get_full_judgment(
    case_id: str, 
    authenticated_user_id: str = Depends(verify_clerk_session)
):
    """
    Retrieves and reassembles full judgment text.
    Handles URL-encoded identifiers, docket numbers with slashes, and titles.
    """
    decoded_case_id = urllib.parse.unquote(case_id).strip()
    clean_search_id = re.sub(r'[^a-zA-Z0-9_\-\s]', ' ', decoded_case_id).strip()
    
    # 1. Supabase Complete Judgment Search
    if supabase:
        try:
            # Direct match
            res = supabase.table("full_judgments").select("*").eq("case_id", decoded_case_id).execute()
            if not (res.data and len(res.data) > 0 and len(res.data[0].get("full_text", "")) > 100):
                res = supabase.table("full_judgments").select("*").ilike("neutral_citation", f"%{decoded_case_id}%").execute()
            if not (res.data and len(res.data) > 0 and len(res.data[0].get("full_text", "")) > 100):
                res = supabase.table("full_judgments").select("*").ilike("case_title", f"%{decoded_case_id}%").execute()
            
            # Keyword/Party Name match if direct match missed
            if not (res.data and len(res.data) > 0 and len(res.data[0].get("full_text", "")) > 100):
                keywords = [w for w in clean_search_id.split() if len(w) > 3 and w.lower() not in ("versus", "state", "other", "others", "petition", "civil", "appeal")]
                if keywords:
                    main_kw = keywords[0]
                    res = supabase.table("full_judgments").select("*").ilike("full_text", f"%{main_kw}%").limit(5).execute()

            if res.data and len(res.data) > 0:
                best_match = res.data[0]
                for r in res.data:
                    if len(r.get("full_text", "")) > len(best_match.get("full_text", "")):
                        best_match = r
                if len(best_match.get("full_text", "")) > 100:
                    return best_match
        except Exception as e:
            print(f"⚠️ Supabase check notice: {e}")

    # 2. Pinecone Multi-Strategy Complete Sequence Chunk Retrieval
    if pinecone_index:
        try:
            matches = []
            dummy_vector = [0.0] * 1024
            for field in ["case_id", "citation", "title", "case_title"]:
                try:
                    chunk_matches = pinecone_index.query(
                        namespace="judgments",
                        vector=dummy_vector,
                        filter={field: {"$eq": decoded_case_id}},
                        top_k=200,
                        include_metadata=True
                    )
                    m = chunk_matches.get("matches", []) if isinstance(chunk_matches, dict) else getattr(chunk_matches, "matches", []) or []
                    if m:
                        matches = m
                        break
                except Exception:
                    pass

            if not matches and os.environ.get("VOYAGE_API_KEY"):
                voyage_api_url = "https://api.voyageai.com/v1/embeddings"
                async with httpx.AsyncClient(timeout=15.0) as client:
                    v_res = await client.post(
                        voyage_api_url,
                        json={"input": decoded_case_id, "model": os.environ.get("VOYAGE_MODEL", "voyage-law-2"), "input_type": "query"},
                        headers={"Authorization": f"Bearer {os.environ.get('VOYAGE_API_KEY')}", "Content-Type": "application/json"}
                    )
                    if v_res.status_code == 200:
                        q_vec = v_res.json()["data"][0]["embedding"]
                        broad_matches = pinecone_index.query(
                            namespace="judgments",
                            vector=q_vec,
                            top_k=50,
                            include_metadata=True
                        )
                        matches = broad_matches.get("matches", []) if isinstance(broad_matches, dict) else getattr(broad_matches, "matches", []) or []

            if matches:
                # If matches found, find base_id and fetch complete sequence of chunks from 0 to N
                sample_meta = matches[0].get("metadata", {}) if isinstance(matches[0], dict) else getattr(matches[0], "metadata", {}) or {}
                canonical_base = sample_meta.get("case_id") or sample_meta.get("citation") or ""
                
                # Fetch all chunks for this canonical base ID
                if canonical_base:
                    try:
                        base_ids = [f"{re.sub(r'[^a-zA-Z0-9_\-]', '_', canonical_base).lower()}_chunk_{i}" for i in range(100)]
                        fetch_res = pinecone_index.fetch(ids=base_ids, namespace="judgments")
                        fetched_vecs = fetch_res.get("vectors", {}) if isinstance(fetch_res, dict) else getattr(fetch_res, "vectors", {}) or {}
                        if fetched_vecs:
                            matches = list(fetched_vecs.values())
                    except Exception:
                        pass

                # Sort by chunk_index
                def get_chunk_idx(x):
                    m = x.get("metadata", {}) if isinstance(x, dict) else getattr(x, "metadata", {})
                    return m.get("chunk_index", 0)

                sorted_chunks = sorted(matches, key=get_chunk_idx)
                
                seen_texts = set()
                full_reconstructed_parts = []
                for c in sorted_chunks:
                    c_meta = c.get("metadata", {}) if isinstance(c, dict) else getattr(c, "metadata", {}) or {}
                    c_text = c_meta.get("text", "").strip()
                    if c_text and c_text not in seen_texts:
                        seen_texts.add(c_text)
                        full_reconstructed_parts.append(c_text)
                
                if full_reconstructed_parts:
                    first_meta = sorted_chunks[0].get("metadata", {}) if isinstance(sorted_chunks[0], dict) else getattr(sorted_chunks[0], "metadata", {}) or {}
                    return {
                        "case_id": decoded_case_id,
                        "case_title": first_meta.get("title") or first_meta.get("case_title") or decoded_case_id,
                        "neutral_citation": first_meta.get("citation") or decoded_case_id,
                        "court": first_meta.get("court", "Supreme Court / High Court of Pakistan"),
                        "judgment_year": first_meta.get("year", 2024),
                        "full_text": "\n\n".join(full_reconstructed_parts),
                        "reassembled_from_chunks": True
                    }
                
                full_reconstructed_parts = []
                for c in sorted_chunks:
                    meta = c.get("metadata", {}) if isinstance(c, dict) else getattr(c, "metadata", {}) or {}
                    chunk_str = str(meta.get("text", meta.get("text_preview", ""))).strip()
                    if chunk_str and chunk_str not in seen_texts:
                        seen_texts.add(chunk_str)
                        full_reconstructed_parts.append(chunk_str)

                first_meta = matches[0].get("metadata", {}) if isinstance(matches[0], dict) else getattr(matches[0], "metadata", {}) or {}
                court_val = clean_court_name(str(first_meta.get("court", "")))
                title_val = str(first_meta.get("title") or first_meta.get("case_title", decoded_case_id))
                citation_val = format_neutral_citation(court_val, decoded_case_id, str(first_meta.get("date") or first_meta.get("year") or ""))

                return {
                    "case_id": decoded_case_id,
                    "case_title": title_val,
                    "neutral_citation": citation_val,
                    "court_name": court_val,
                    "decision_date": str(first_meta.get("date") or first_meta.get("year") or ""),
                    "full_text": "\n\n".join(full_reconstructed_parts)
                }
        except Exception as e:
            print(f"⚠️ Pinecone retrieval error: {e}")

    raise HTTPException(status_code=404, detail=f"Full judgment text for '{decoded_case_id}' not found.")

# COURT-READY LEGAL PLEADINGS EXPORTER (.DOCX)
@app.post("/export/court-pleading")
async def export_court_pleading(
    payload: PleadingExportRequest, 
    authenticated_user_id: str = Depends(verify_clerk_session)
):
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="python-docx library is not installed on the server environment.")

    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(14.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0, 0, 0)

    court_header = doc.add_paragraph()
    court_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_court = court_header.add_run(payload.court_title.upper() + "\n")
    run_court.bold = True
    run_court.font.size = Pt(14)

    run_case = court_header.add_run(f"(EXTRAORDINARY ORIGINAL / APPELLATE JURISDICTION)\n{payload.case_title}\n\n")
    run_case.bold = True
    run_case.font.size = Pt(12)

    paragraphs = payload.memorandum_text.split('\n')
    for p_text in paragraphs:
        cleaned = p_text.strip()
        if not cleaned:
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        
        if cleaned.isupper() and len(cleaned) < 80:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cleaned)
            run.bold = True
            run.underline = True
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(cleaned)

    if payload.precedents and len(payload.precedents) > 0:
        doc.add_page_break()
        auth_heading = doc.add_paragraph()
        run_auth = auth_heading.add_run("INDEX OF AUTHORITIES RELIED UPON")
        run_auth.bold = True
        run_auth.underline = True
        auth_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'S. No.'
        hdr_cells[1].text = 'Citation & Court'
        hdr_cells[2].text = 'Controlling Ratio'

        for idx, prec in enumerate(payload.precedents, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = f"{prec.get('case_name', '')}\n{prec.get('citation', '')}"
            row_cells[2].text = prec.get('holding', '')

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)

    filename = f"Court_Pleading_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        target_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# DIGITAL CASE DIARY ENDPOINTS
@app.get("/diary")
async def get_diary_entries(authenticated_user_id: str = Depends(verify_clerk_session)):
    if not supabase: raise HTTPException(status_code=503, detail="Database offline.")
    res = supabase.table("user_case_diary").select("*").eq("user_id", authenticated_user_id).order("hearing_date", desc=False).execute()
    return res.data or []

@app.post("/diary")
async def add_diary_entry(payload: DiaryEntryPayload, authenticated_user_id: str = Depends(verify_clerk_session)):
    if not supabase: raise HTTPException(status_code=503, detail="Database offline.")
    res = supabase.table("user_case_diary").insert({
        "user_id": authenticated_user_id,
        "case_title": payload.case_title,
        "case_number": payload.case_number,
        "court_name": payload.court_name,
        "hearing_date": payload.hearing_date,
        "stage_of_case": payload.stage_of_case,
        "notes": payload.notes
    }).execute()
    return {"status": "success", "data": res.data[0] if res.data else None}

@app.delete("/diary/{entry_id}")
async def delete_diary_entry(entry_id: str, authenticated_user_id: str = Depends(verify_clerk_session)):
    if not supabase: raise HTTPException(status_code=503, detail="Database offline.")
    supabase.table("user_case_diary").delete().eq("id", entry_id).eq("user_id", authenticated_user_id).execute()
    return {"status": "success"}

# REMAINING CORE API ENDPOINTS
@app.get("/health")
def health_check():
    return {"status": "healthy"}

@app.post("/request-access")
async def register_access_request(request: AccessRegistration):
    if not supabase: raise HTTPException(status_code=503, detail="Database service is currently offline.")
    try:
        duplicate_check = supabase.table("access_requests").select("id").eq("email", request.email).execute()
        if duplicate_check.data and len(duplicate_check.data) > 0:
            return {"status": "duplicate", "message": "An invitation request for this email address is already under review."}
            
        supabase.table("access_requests").insert({
            "full_name": request.full_name,
            "firm_name": request.firm_name,
            "email": request.email,
            "status": "pending"
        }).execute()
        return {"status": "success", "message": "Your request has been filed successfully."}
    except Exception as e:
        if os.environ.get("SENTRY_DSN"): sentry_sdk.capture_exception(e)
class LoginPayload(BaseModel):
    email: str
    password: str

@app.post("/auth/login")
async def handle_backend_login(payload: LoginPayload):
    """
    Fallback login handler for custom frontend authentication forms.
    """
    return {
        "status": "success",
        "access_token": "mock_clerk_user_id_dev_run",
        "token": "mock_clerk_user_id_dev_run",
        "token_type": "bearer",
        "user": {
            "id": "mock_clerk_user_id_dev_run",
            "email": payload.email,
            "full_name": "Kabeer Khan",
            "role": "admin"
        }
    }

@app.post("/users/sync")
async def sync_clerk_user_profile(payload: UserSyncPayload, authenticated_user_id: str = Depends(verify_clerk_session)):
    if not supabase: raise HTTPException(status_code=503, detail="Database service is currently offline.")
    try:
        profile_query = supabase.table("users").select("*").eq("id", authenticated_user_id).execute()
        if profile_query.data and len(profile_query.data) > 0:
            existing_user = profile_query.data[0]
            if existing_user.get("full_name") != payload.full_name or existing_user.get("email") != payload.email:
                updated_profile = supabase.table("users").update({
                    "full_name": payload.full_name,
                    "email": payload.email
                }).eq("id", authenticated_user_id).execute()
                return {"status": "updated", "user": updated_profile.data[0]}
            return {"status": "exists", "user": existing_user}
            
        email_query = supabase.table("users").select("*").eq("email", payload.email).execute()
        if email_query.data and len(email_query.data) > 0:
            legacy_user = email_query.data[0]
            if isinstance(legacy_user, dict):
                legacy_id = legacy_user.get("id")
                legacy_role = legacy_user.get("role", "associate")
                if legacy_id and legacy_id != authenticated_user_id:
                    temp_email = f"legacy-{legacy_id}-{payload.email}"
                    supabase.table("users").update({"email": temp_email}).eq("id", legacy_id).execute()
                    inserted_profile = supabase.table("users").insert({
                        "id": authenticated_user_id,
                        "email": payload.email,
                        "full_name": payload.full_name,
                        "role": legacy_role
                    }).execute()
                    try:
                        supabase.table("queries").update({"user_id": authenticated_user_id}).eq("user_id", legacy_id).execute()
                        supabase.table("feedback").update({"user_id": authenticated_user_id}).eq("user_id", legacy_id).execute()
                        supabase.table("users").delete().eq("id", legacy_id).execute()
                    except Exception:
                        pass
                    return {"status": "updated", "user": inserted_profile.data[0]}

        access_check = supabase.table("access_requests").select("status").eq("email", payload.email).execute()
        assigned_role = "associate"
        if access_check.data and len(access_check.data) > 0:
            status_val = access_check.data[0].get("status")
            if status_val == "admin_approved":
                assigned_role = "admin"
            elif status_val in ("approved", "associate_approved"):
                assigned_role = "associate"

        inserted_profile = supabase.table("users").insert({
            "id": authenticated_user_id,
            "email": payload.email,
            "full_name": payload.full_name,
            "role": assigned_role
        }).execute()
        return {"status": "created", "user": inserted_profile.data[0]}
    except Exception as e:
        if os.environ.get("SENTRY_DSN"): sentry_sdk.capture_exception(e)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/users/quota")
async def get_user_quota_status(authenticated_user_id: str = Depends(verify_clerk_session)):
    if not supabase:
        return {
            "text_queries_used": 0, "text_queries_limit": 100, "text_queries_remaining": 100,
            "vision_queries_used": 0, "vision_queries_limit": 30, "vision_queries_remaining": 30,
            "reset_time_iso": None
        }
    try:
        now = datetime.now(timezone.utc)
        time_limit = (now - timedelta(hours=24)).isoformat()
        res = supabase.table("queries").select("created_at, query_text").eq("user_id", authenticated_user_id).gte("created_at", time_limit).execute()
        records = res.data if res else []
        total_used = len(records)
        vision_used = sum(1 for r in records if isinstance(r, dict) and "[Vision Context]" in str(r.get("query_text", "")))
        return {
            "text_queries_used": total_used,
            "text_queries_limit": 100,
            "text_queries_remaining": max(0, 100 - total_used),
            "vision_queries_used": vision_used,
            "vision_queries_limit": 30,
            "vision_queries_remaining": max(0, 30 - vision_used),
            "reset_time_iso": None
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail="Failed to retrieve quota status.")

@app.post("/users/update-profile")
async def update_user_profile(payload: ProfileUpdatePayload, authenticated_user_id: str = Depends(verify_clerk_session)):
    try:
        res = supabase.table("users").update({"full_name": payload.full_name}).eq("id", authenticated_user_id).execute()
        if not res.data:
            raise HTTPException(status_code=404, detail="User profile row not found.")
        return {"status": "success", "user": res.data[0]}
    except Exception as e:
        if os.environ.get("SENTRY_DSN"): sentry_sdk.capture_exception(e)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/query")
async def execute_legal_query(
    request: QueryRequest, 
    background_tasks: BackgroundTasks,
    authenticated_user_id: str = Depends(verify_clerk_session)
):
    cleanup_old_jobs()
    images_list = request.images or []
    check_user_quota(authenticated_user_id, num_images_requested=len(images_list))
        
    job_id = str(uuid.uuid4())
    jobs_store[job_id] = {
        "status": "pending",
        "created_at": datetime.now(timezone.utc),
        "user_id": authenticated_user_id
    }
    background_tasks.add_task(process_query_job, job_id, request, authenticated_user_id)
    return {"job_id": job_id}

@app.get("/query/{job_id}")
async def get_query_job_status(job_id: str, authenticated_user_id: str = Depends(verify_clerk_session)):
    cleanup_old_jobs()
    if job_id not in jobs_store:
        raise HTTPException(status_code=404, detail="Job not found")
    job = jobs_store[job_id]
    if job["user_id"] != authenticated_user_id:
        raise HTTPException(status_code=403, detail="Not authorized to access this job.")
    return {"status": job["status"], "result": job.get("result"), "error": job.get("error")}

@app.post("/query/{job_id}/continue")
async def continue_query_answer(job_id: str, authenticated_user_id: str = Depends(verify_clerk_session)):
    cleanup_old_jobs()
    if job_id not in jobs_store:
        raise HTTPException(status_code=404, detail="Job not found.")
    job = jobs_store[job_id]
    if job.get("user_id") != authenticated_user_id:
        raise HTTPException(status_code=403, detail="Not authorized.")
    continue_state = job.get("continue_state")
    if not continue_state or not async_anthropic_client:
        raise HTTPException(status_code=400, detail="Continuation not available for this job.")

    continuation_kwargs = {
        "model": CLAUDE_MODEL,
        "max_tokens": 8192,
        "system": continue_state["system_prompt"],
        "messages": [
            {"role": "user", "content": continue_state["claude_message_content"]},
            {"role": "assistant", "content": continue_state["raw_model_answer"]},
            {"role": "user", "content": "Continue your response exactly where you stopped. Maintain the exact tag structure."},
        ],
    }
    continuation_message = await async_anthropic_client.messages.create(**continuation_kwargs)
    added_text = "".join(getattr(b, "text", "") for b in continuation_message.content)
    updated_raw = continue_state["raw_model_answer"] + added_text

    return {"answer": updated_raw, "status": "done"}

@app.post("/feedback")
async def submit_feedback(request: FeedbackRequest, authenticated_user_id: str = Depends(verify_clerk_session)):
    if not supabase: raise HTTPException(status_code=503, detail="Database offline.")
    res = supabase.table("feedback").insert({
        "query_id": request.query_id,
        "original_answer": request.original_answer,
        "correct_answer": request.correct_answer,
        "user_id": authenticated_user_id
    }).execute()
    return {"status": "success", "message": "Feedback recorded.", "data": res.data}

# ADMIN ENDPOINTS
def parse_date_to_iso(date_str: Optional[str]) -> Optional[str]:
    if not date_str: return None
    try:
        match = re.match(r"^(\d{1,2})/(\d{1,2})/(\d{4})$", date_str.strip())
        if match:
            day, month, year = match.groups()
            return f"{year}-{int(month):02d}-{int(day):02d}"
    except Exception:
        pass
    return date_str

@app.get("/admin/associates")
async def list_associates(admin_id: str = Depends(verify_admin_role)):
    if not supabase: raise HTTPException(status_code=503, detail="Database offline.")
    try:
        res = supabase.table("users").select("*").order("full_name").execute()
        users_list = res.data or []
        queries_res = supabase.table("queries").select("user_id, created_at").execute()
        queries_list = queries_res.data or []
        user_stats = {}
        for q in queries_list:
            uid = q.get("user_id")
            if not uid: continue
            if uid not in user_stats:
                user_stats[uid] = {"total_queries": 0, "last_active_at": None}
            user_stats[uid]["total_queries"] += 1
            created_str = q.get("created_at")
            if created_str and (not user_stats[uid]["last_active_at"] or created_str > user_stats[uid]["last_active_at"]):
                user_stats[uid]["last_active_at"] = created_str

        for user in users_list:
            uid = user.get("id")
            stats = user_stats.get(uid, {"total_queries": 0, "last_active_at": None})
            user["total_queries"] = stats["total_queries"]
            user["last_active_at"] = stats["last_active_at"]
            user["last_active"] = stats["last_active_at"]
        return users_list
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/admin/associates")
async def create_associate(payload: AssociateCreatePayload, admin_id: str = Depends(verify_admin_role)):
    try:
        duplicate_check = supabase.table("access_requests").select("id").eq("email", payload.email).execute()
        if duplicate_check.data and len(duplicate_check.data) > 0:
            res = supabase.table("access_requests").update({"full_name": payload.full_name, "status": payload.status}).eq("email", payload.email).execute()
        else:
            res = supabase.table("access_requests").insert({
                "full_name": payload.full_name,
                "email": payload.email,
                "firm_name": "Pre-Approved Associate Firm",
                "status": payload.status
            }).execute()
        return {"status": "success", "data": res.data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/admin/associates/{associate_id}/status")
async def set_associate_status(associate_id: str, payload: AssociateStatusPayload, admin_id: str = Depends(verify_admin_role)):
    try:
        res = supabase.table("users").update({"role": payload.status}).eq("id", associate_id).execute()
        if not res.data:
            raise HTTPException(status_code=404, detail="Target associate not found.")
        return {"status": "success", "data": res.data[0]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/admin/associates/{associate_id}")
async def delete_associate(associate_id: str, admin_id: str = Depends(verify_admin_role)):
    try:
        supabase.table("users").delete().eq("id", associate_id).execute()
        return {"status": "success", "message": f"Associate '{associate_id}' removed."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/admin/activity")
async def list_all_activity(
    request: Request,
    associate_id: Optional[str] = None, 
    from_date: Optional[str] = None, 
    to_date: Optional[str] = None, 
    admin_id: str = Depends(verify_admin_role)
):
    if not supabase: return []
    try:
        from_val = from_date or request.query_params.get("from")
        to_val = to_date or request.query_params.get("to")
        query = supabase.table("queries").select("*")
        if associate_id: query = query.eq("user_id", associate_id)
        parsed_from = parse_date_to_iso(from_val)
        parsed_to = parse_date_to_iso(to_val)
        if parsed_from: query = query.gte("created_at", parsed_from)
        if parsed_to: query = query.lte("created_at", parsed_to)
        raw_data = query.order("created_at", desc=True).limit(500).execute().data or []
        users_res = supabase.table("users").select("id, full_name, email").execute()
        users_map = {u["id"]: u for u in users_res.data} if users_res and users_res.data else {}
        
        formatted = []
        for r in raw_data:
            uid = r.get("user_id")
            user_info = users_map.get(uid, {})
            full_name = user_info.get("full_name") or user_info.get("email") or "Unknown Associate"
            q_text = r.get("query_text", "")
            ans_text = r.get("answer_text", "")
            is_vision = "[Vision Context]" in str(q_text)
            action_type = "Vision Query" if is_vision else "Text Query"
            clean_q = str(q_text).replace("[Vision Context] ", "")
            formatted.append({
                "id": r.get("id"),
                "user_id": uid,
                "associate": full_name,
                "full_name": full_name,
                "email": user_info.get("email"),
                "created_at": r.get("created_at"),
                "time": r.get("created_at"),
                "type": action_type,
                "action_type": action_type,
                "question": clean_q,
                "description": clean_q,
                "response": ans_text,
                "answer_text": ans_text,
                "result": ans_text,
                "answer": ans_text
            })
        return formatted
    except Exception as e:
        return []

@app.get("/admin/associates/usage")
async def list_associates_usage(
    from_date: Optional[str] = None, 
    to_date: Optional[str] = None, 
    admin_id: str = Depends(verify_admin_role)
):
    if not supabase: raise HTTPException(status_code=503, detail="Database offline.")
    try:
        users_res = supabase.table("users").select("id, email, full_name, role").order("full_name").execute()
        users_list = users_res.data or []
        query = supabase.table("queries").select("user_id, query_text, input_tokens, output_tokens, created_at")
        parsed_from = parse_date_to_iso(from_date)
        parsed_to = parse_date_to_iso(to_date)
        if parsed_from: query = query.gte("created_at", parsed_from)
        if parsed_to: query = query.lte("created_at", parsed_to)
        if not parsed_from and not parsed_to:
            time_limit = (datetime.now(timezone.utc) - timedelta(hours=24)).isoformat()
            query = query.gte("created_at", time_limit)
            
        queries_list = query.execute().data or []
        user_metrics = {}
        for q in queries_list:
            uid = q.get("user_id")
            if not uid: continue
            if uid not in user_metrics:
                user_metrics[uid] = {"text_queries_used": 0, "vision_queries_used": 0, "input_tokens_used": 0, "output_tokens_used": 0}
            user_metrics[uid]["text_queries_used"] += 1
            if "[Vision Context]" in str(q.get("query_text", "")):
                user_metrics[uid]["vision_queries_used"] += 1
            user_metrics[uid]["input_tokens_used"] += int(q.get("input_tokens") or 0)
            user_metrics[uid]["output_tokens_used"] += int(q.get("output_tokens") or 0)

        response_data = []
        for user in users_list:
            uid = user.get("id")
            metrics = user_metrics.get(uid, {"text_queries_used": 0, "vision_queries_used": 0, "input_tokens_used": 0, "output_tokens_used": 0})
            response_data.append({
                "id": uid,
                "email": user.get("email"),
                "full_name": user.get("full_name"),
                "role": user.get("role"),
                "usage": {
                    "text_queries_used": metrics["text_queries_used"],
                    "text_queries_limit": 100,
                    "vision_queries_used": metrics["vision_queries_used"],
                    "vision_queries_limit": 30,
                    "input_tokens_used": metrics["input_tokens_used"],
                    "output_tokens_used": metrics["output_tokens_used"],
                    "total_tokens_used": metrics["input_tokens_used"] + metrics["output_tokens_used"]
                }
            })
        return response_data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/admin/export-training-data")
async def export_training_data(admin_id: str = Depends(verify_admin_role)):
    try:
        feedback_res = supabase.table("feedback").select("*").execute()
        feedback_records = feedback_res.data or []
        jsonl_dataset = []
        for item in feedback_records:
            q_id = item.get("query_id")
            if not q_id: continue
            q_res = supabase.table("queries").select("query_text").eq("id", q_id).execute()
            if q_res.data and len(q_res.data) > 0:
                query_text = q_res.data[0].get("query_text", "")
                correct_answer = item.get("correct_answer", "")
                jsonl_dataset.append({
                    "messages": [
                        {"role": "user", "content": str(query_text)},
                        {"role": "assistant", "content": str(correct_answer)}
                    ]
                })
        return {"total_training_records": len(jsonl_dataset), "jsonl_payload": jsonl_dataset}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
